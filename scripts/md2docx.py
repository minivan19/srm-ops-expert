#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word脚本
用法：
  python scripts/md2docx.py 客户名
  python scripts/md2docx.py 虎牙 --year 2025
"""

import os
import re
import glob
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

CLIENT_DATA_DIR = "/Users/limingheng/AI/client-data"


def find_latest_md(client_name, year):
    """查找最新版本的运维报告MD文件"""
    pattern = os.path.join(CLIENT_DATA_DIR, client_name, f"{client_name}_{year}_运维报告_V*.md")
    md_files = glob.glob(pattern)
    if not md_files:
        return None, None
    # 按版本号排序
    def extract_version(f):
        match = re.search(r'V(\d+)', f)
        return int(match.group(1)) if match else 0
    md_files.sort(key=extract_version, reverse=True)
    latest = md_files[0]
    doc_file = latest.replace('.md', '.docx')
    return latest, doc_file


def read_md(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, font_size=None):
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    if font_size:
        p.runs[0].font.size = Pt(font_size)
    return p


def add_table_from_markdown(doc, md_content):
    """从markdown表格转换"""
    lines = md_content.strip().split('\n')
    table = None

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'\|[\s\-:|]+\|', line):
            continue

        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]

            if table is None:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                header_row = table.rows[0]
                for i, cell in enumerate(cells):
                    header_row.cells[i].text = cell
            else:
                row = table.add_row()
                for i, cell in enumerate(cells):
                    if i < len(row.cells):
                        row.cells[i].text = cell

    return table is not None


def parse_and_convert(md_content, doc):
    """解析markdown并添加到doc"""
    lines = md_content.split('\n')
    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith('|'):
            table_buffer.append(line)
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith('|'):
                add_table_from_markdown(doc, '\n'.join(table_buffer))
                table_buffer = []
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 3)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph()
            while '**' in text:
                start = text.find('**')
                end = text.find('**', start + 2)
                if end == -1:
                    break
                before = text[:start]
                bold_text = text[start+2:end]
                after = text[end+2:]

                if before:
                    p.add_run(before)
                p.add_run(bold_text).bold = True
                text = after

            if text:
                p.add_run(text)
        else:
            text = line
            p = doc.add_paragraph()
            while '**' in text:
                start = text.find('**')
                end = text.find('**', start + 2)
                if end == -1:
                    break
                before = text[:start]
                bold_text = text[start+2:end]
                after = text[end+2:]

                if before:
                    p.add_run(before)
                p.add_run(bold_text).bold = True
                text = after

            if text:
                p.add_run(text)

        i += 1


def main():
    parser = argparse.ArgumentParser(description="Markdown转Word")
    parser.add_argument("client_name", help="客户名称")
    parser.add_argument("--year", type=int, default=None, help="年份（不指定则默认上一自然年）")
    args = parser.parse_args()

    client_name = args.client_name
    year = args.year if args.year else datetime.now().year - 1

    MD_FILE, DOC_FILE = find_latest_md(client_name, year)
    if MD_FILE is None:
        print(f"错误：找不到 {client_name} {year} 年运维报告MD文件")
        print(f"请先运行：python generate_report_v2.py {client_name}")
        return

    print(f"读取: {MD_FILE}")
    md_content = read_md(MD_FILE)
    print(f"  - 字符数: {len(md_content)}")

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)

    print("转换中...")
    parse_and_convert(md_content, doc)

    print(f"保存: {DOC_FILE}")
    doc.save(DOC_FILE)
    print("完成!")


if __name__ == "__main__":
    main()
