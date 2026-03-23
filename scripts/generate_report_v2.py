#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SRM运维报告生成脚本
自动调用LLM生成完整报告
用法：
  python scripts/generate_report_v2.py 客户名
"""

import os
import sys
import glob
import requests
import pandas as pd
import time
import argparse
from datetime import datetime

# ============== 配置 ==============
RAW_DATA_ROOT = "/Users/limingheng/AI/client-data/raw/客户档案"
CLIENT_DATA_DIR = "/Users/limingheng/AI/client-data"
WORKSPACE_MEDIA_DIR = "/Users/limingheng/.openclaw/workspace/media"

# API配置
API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-340ed7819c2346508c0a46a80df85999")

# LLM配置
MODEL = "deepseek-chat"
TEMPERATURE = 0.3

# 全局路径（由 build_paths 填充）
RAW_DATA_DIR = None
MODULE_DATA_FILE = None
OUTPUT_DIR = None
CLIENT_NAME = None
REPORT_YEAR = None


def build_paths(client_name, year):
    """根据客户名构建路径"""
    global RAW_DATA_DIR, MODULE_DATA_FILE, OUTPUT_DIR, CLIENT_NAME, REPORT_YEAR
    CLIENT_NAME = client_name
    REPORT_YEAR = year
    raw_dir = os.path.join(RAW_DATA_ROOT, client_name, "运维工单")
    out_dir = os.path.join(CLIENT_DATA_DIR, client_name)
    os.makedirs(out_dir, exist_ok=True)
    module_file = os.path.join(out_dir, f"{client_name}_{year}_模块工单数据.txt")
    RAW_DATA_DIR = raw_dir
    MODULE_DATA_FILE = module_file
    OUTPUT_DIR = out_dir


def call_llm(prompt, temperature=TEMPERATURE, max_retries=5, retry_delay=5.0,
              max_tokens=800):
    """调用DeepSeek LLM（带重试机制，非streaming模式）"""
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens
    }
    last_error = None
    for attempt in range(1, max_retries + 1):
        try:
            # 使用 stream=False，读取完整响应后再解析
            response = requests.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=(30, 240),
                stream=False
            )
            if response.status_code == 200:
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    return result["choices"][0]["message"]["content"]
                else:
                    last_error = f"无效响应格式: {result}"
            else:
                last_error = f"HTTP {response.status_code}: {response.text[:200]}"
        except Exception as e:
            last_error = str(e)
        if attempt < max_retries:
            print(f"    [LLM重试 {attempt}/{max_retries}，{retry_delay:.0f}s后...]")
            time.sleep(retry_delay)
            retry_delay *= 2
    raise RuntimeError(f"LLM调用失败（已重试{max_retries}次）: {last_error}")


def load_module_data():
    """读取模块工单数据"""
    if not os.path.exists(MODULE_DATA_FILE):
        print(f"错误: 数据文件不存在 {MODULE_DATA_FILE}")
        print("请先运行 extract_module_data.py")
        sys.exit(1)
    
    with open(MODULE_DATA_FILE, 'r', encoding='utf-8') as f:
        return f.read()


def load_raw_data():
    """读取原始Excel数据用于统计（仅指定年份）"""
    files = glob.glob(os.path.join(RAW_DATA_DIR, "*.xlsx"))
    files = [f for f in files if str(REPORT_YEAR) in os.path.basename(f)]
    if not files:
        print(f"错误: 未找到{REPORT_YEAR}年Excel文件 {RAW_DATA_DIR}")
        sys.exit(1)

    df = pd.concat([pd.read_excel(f) for f in files], ignore_index=True)
    return df


def get_statistics(df):
    """获取统计数据"""
    # 月度分布 - 表格格式
    df['创建时间'] = pd.to_datetime(df['创建时间'], errors='coerce')
    df['月份'] = df['创建时间'].dt.month
    monthly = df['月份'].value_counts().sort_index()
    monthly_rows = []
    for m in range(1, 13):
        cnt = int(monthly.get(m, 0))
        monthly_rows.append(f"| {m}月 | {cnt} |")
    monthly_table = "\n".join(["| 月份 | 工单数 |"] + ["|------|--------|"] + monthly_rows)
    
    # 分类统计 - 表格格式
    category_counts = df['分类'].value_counts()
    total = len(df)
    category_rows = []
    for i, (cat, cnt) in enumerate(category_counts.items(), 1):
        pct = cnt / total * 100
        category_rows.append(f"| {i} | {cat} | {cnt} | {pct:.1f}% |")
    category_table = "\n".join(["| 序号 | 分类名称 | 工单数 | 占比 |"] + ["|------|---------|--------|-------|"] + category_rows)
    
    # 模块统计 - 表格格式
    module_counts = df['模块'].value_counts()
    module_rows = []
    for i, (mod, cnt) in enumerate(module_counts.items(), 1):
        pct = cnt / total * 100
        module_rows.append(f"| {i} | {mod} | {cnt} | {pct:.1f}% |")
    module_table = "\n".join(["| 序号 | 系统模块 | 工单数 | 占比 |"] + ["|------|---------|--------|-------|"] + module_rows)
    
    # SLA和满意度
    sla_rate = df['SLA是否达标'].value_counts(normalize=True).get('达标', 0) * 100
    product_sat = df['产品满意度'].mean() if '产品满意度' in df.columns else 5.0
    service_sat = df['服务满意度'].mean() if '服务满意度' in df.columns else 5.0
    
    return {
        'monthly': monthly_table,
        'category': category_table,
        'module': module_table,
        'total': len(df),
        'sla': f"{sla_rate:.1f}%",
        'product_sat': f"{product_sat:.1f}",
        'service_sat': f"{service_sat:.1f}"
    }


def batch_0_trend_sla(stats):
    """第0批：趋势+SLA分析"""
    prompt = f"""## 任务：基于统计数据生成趋势分析和SLA分析

### 数据
- 月度分布：{stats['monthly']}
- SLA达标率：{stats['sla']}
- 产品满意度：{stats['product_sat']}分
- 服务满意度：{stats['service_sat']}分

### 输出格式
### 1.4 趋势分析
**整体趋势判断** [简要描述]
**趋势预测与建议** 1. 2. 3.

### 1.5 SLA与满意度分析
**核心指标表现**
- SLA达标率：{stats['sla']} → [评价]
- 产品满意度：{stats['product_sat']}分 → [评价]
- 服务满意度：{stats['service_sat']}分 → [评价]"""

    return call_llm(prompt)


def batch_1_classification_module(module_data, stats):
    """第1批：分类维度分析（循环调用，每个分类单独处理）"""
    import re

    # 第1部分：从 stats 提取 >=10 单的分类
    cat_lines = [l for l in stats['category'].strip().split('\n') if '|' in l and '分类' not in l]
    valid_cats = []  # [(分类名, 工单数)]
    for line in cat_lines:
        parts = line.strip('| ').split('|')
        if len(parts) >= 2:
            try:
                cnt = int(parts[-1].strip())
                name = '|'.join(parts[:-1]).strip()
                if cnt >= 10:
                    valid_cats.append((name, cnt))
            except ValueError:
                pass

    # 第2部分：从 module_data 按分类聚合工单
    # module_data 格式：### 模块名 (N单)\n**工单号**: ...\n**分类**: xxx\n...
    tickets_by_cat = {}  # {分类名: [ticket_block]}
    sections = module_data.split("### ")
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
        lines = sec.split('\n')
        # 解析该模块下的工单（每个工单以 **工单号** 开头）
        current_ticket = []
        for line in lines:
            if line.startswith('**工单号**'):
                # 保存上一个工单
                if current_ticket:
                    ticket_text = '\n'.join(current_ticket)
                    # 提取分类
                    cat_m = re.search(r'\*\*分类\*\*:\s*(.+)', ticket_text)
                    if cat_m:
                        cat_name = cat_m.group(1).strip()
                        if cat_name not in tickets_by_cat:
                            tickets_by_cat[cat_name] = []
                        tickets_by_cat[cat_name].append(ticket_text)
                    current_ticket = [line]
                else:
                    current_ticket = [line]
            else:
                if current_ticket:
                    current_ticket.append(line)
        # 处理最后一张工单
        if current_ticket:
            ticket_text = '\n'.join(current_ticket)
            cat_m = re.search(r'\*\*分类\*\*:\s*(.+)', ticket_text)
            if cat_m:
                cat_name = cat_m.group(1).strip()
                if cat_name not in tickets_by_cat:
                    tickets_by_cat[cat_name] = []
                tickets_by_cat[cat_name].append(ticket_text)

    # 第3部分：循环调用 LLM，每个分类单独处理
    results = []
    for idx, (cat_name, count) in enumerate(valid_cats, 1):
        tickets = tickets_by_cat.get(cat_name, [])
        if not tickets:
            results.append(f"### 2.{idx}. {cat_name} - {count}单\n（无工单数据）")
            continue
        section_data = '\n\n---\n\n'.join(tickets[:50])  # 最多50条
        num_problems = 3 if count >= 30 else (2 if count >= 15 else 1)
        time.sleep(10)

        prompt = f"""## 任务：基于以下工单数据，按分类分析主要问题

### 数据（{cat_name}，共{count}单）
{section_data}

### 重要约束
1. 只基于工单数据，不要推测
2. 解决方案有则提取，无则写"待完善"
3. 参考工单必须真实存在
4. 禁止空洞词汇

### 输出格式
### 2.{idx}. {cat_name} - {count}单，建议{num_problems}个问题
#### 问题1：[问题主题]
**问题描述**：...
**原因分析**：1. ...
**解决方案**：1. ...
**参考工单**：I-xxx

#### 问题2：[问题主题]
（如需生成{num_problems}个问题则继续）

注意：每个问题前必须加`#### 问题[N]`子标题以示区分"""

        try:
            result = call_llm(prompt, max_tokens=3000)
            results.append(result)
            print(f"  [{idx}/{len(valid_cats)}] {cat_name} ({count}单) OK")
        except Exception as e:
            print(f"  [{idx}/{len(valid_cats)}] {cat_name} ({count}单) 失败: {e}")
            results.append(f"### 2.{idx}. {cat_name} - {count}单\n（分析失败）")

    result_cat = '\n\n'.join(results)

    # 第2部分：模块维度分析（循环调用，只处理 >=10 单的模块）
    sections = module_data.split("### ")
    valid_sections = []
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
        lines = sec.split("\n")
        module_name = lines[0].strip()
        if not module_name:
            continue
        m = re.search(r'\((\d+)单\)', module_name)
        count = int(m.group(1)) if m else 0
        if count >= 10:
            valid_sections.append((module_name, count, "\n".join(lines[:100])))
    results = []
    for idx, (module_name, count, section_data) in enumerate(valid_sections, 1):
        num_problems = 3 if count >= 30 else (2 if count >= 15 else 1)
        time.sleep(10)

        prompt = f"""## 任务：基于以下工单数据，分析主要问题

### 数据（{module_name}，共{count}单）
{section_data}

### 重要约束
1. 只基于工单数据，不要推测
2. 解决方案有则提取，无则写"待完善"
3. 参考工单必须真实存在
4. 禁止空洞词汇

### 输出格式
### 3.{idx}. {module_name}
#### 问题1：[问题主题]
**问题描述**：...
**原因分析**：1. ...
**解决方案**：1. ...
**参考工单**：I-xxx

#### 问题2：[问题主题]
（如需生成{num_problems}个问题则继续）

注意：每个问题前必须加`#### 问题[N]`子标题以示区分"""

        try:
            result = call_llm(prompt, max_tokens=2500)
            results.append(result)
            print(f"  [{idx}/{len(valid_sections)}] {module_name} OK")
        except Exception as e:
            print(f"  [{idx}/{len(valid_sections)}] {module_name} 失败: {e}")
            results.append(f"### 3.{idx}. {module_name}\n（数据不足，无法分析）")

    result_mod = "\n\n".join(results)
    return result_cat, result_mod


def batch_2_faq(module_data):
    """第2批：FAQ"""
    prompt = f"""## 任务：从业务咨询类工单生成高频FAQ（面向最终用户）

### 重要约束
1. 基于工单的"标题"、"描述"、"解决方案"、"根本原因"四个字段综合分析
2. 解决方案定位：面向最终用户（业务人员），不是技术人员
3. 解决方案类型限制：
   - ✅ 可以是：操作指导（如：点击XX按钮→填写XX→提交）
   - ✅ 可以是：答疑解释（如：这是因为...）
   - ✅ 可以是：联系系统管理员（如：请联系管理员处理"XX问题"）
   - ❌ 不能是：配置建议（如：进入系统设置→配置XX）
   - ❌ 不能是：实施优化建议（如：建议优化XX流程）
4. 禁止使用"优化"、"完善"、"配置"、"调整"等词汇
5. 问题场景需要聚类提炼，不强求使用原文
6. 参考工单必须与问题描述直接相关，禁止幻觉

### 数据字段
- 标题：工单标题
- 描述：详细描述
- 解决方案：处理方案
- 根本原因：根因分析

### 数据
{module_data}

### 输出格式
### Q1：{{基于多个工单标题聚类提炼}}
**问题场景** [综合标题+描述提炼，使用用户能理解的语言]
**解决方案** 
1. [操作指导：如果问题可以通过操作解决，给出具体步骤]
2. [或答疑解释：如果无法操作解决，解释原因]
3. [或联系管理员：如果需要系统配置，写"请联系系统管理员处理"]
**参考工单** - I-xxxxx"""

    return call_llm(prompt, max_tokens=2000)


def batch_3_summary(batch1_result, batch2_result, stats):
    """第3批：总结"""
    prompt = f"""## 任务：基于前期分析结果生成总结与改进建议

### 第1批分析结果
{batch1_result[:2000]}

### 第2批FAQ结果
{batch2_result[:1000]}

### 基础数据
- 总工单：{stats['total']}
- SLA达标率：{stats['sla']}
- 满意度：{stats['product_sat']}分

### 输出格式
### 5.1 运维整体表现
### 5.2 主要问题与风险
### 5.3 改进建议"""

    return call_llm(prompt, max_tokens=2000)


def generate_report():
    """生成完整报告"""
    print("=" * 50)
    print(f"SRM Report Generator - {CLIENT_NAME}")
    print("=" * 50)

    # 1. 加载数据
    print("\n[1/5] Loading data...")
    module_data = load_module_data()
    df = load_raw_data()
    stats = get_statistics(df)
    print(f"  Total tickets: {stats['total']}")
    print(f"  SLA rate: {stats['sla']}")

    # 2. 第0批：趋势+SLA
    print("\n[2/5] Batch 0: Trend+SLA...")
    result_0 = batch_0_trend_sla(stats)
    print("  [OK]")

    # 3. 第1批：分类+模块
    print("\n[3/5] Batch 1: Classification+Module...")
    result_1, result_2 = batch_1_classification_module(module_data, stats)
    print("  [OK]")

    # 4. 第2批：FAQ
    print("\n[4/5] Batch 2: FAQ...")
    result_3 = batch_2_faq(module_data)
    print("  [OK]")

    # 5. 第3批：总结
    print("\n[5/5] Batch 3: Summary...")
    result_4 = batch_3_summary(result_1, result_2, stats)
    print("  [OK]")

    # 整合报告
    report = f"""# {CLIENT_NAME}{REPORT_YEAR}年度SRM运维分析报告

## 报告信息
| 项目 | 内容 |
|------|------|
| 报告周期 | {REPORT_YEAR}年度 |
| 客户名称 | {CLIENT_NAME} |
| 报告生成时间 | {datetime.now().strftime('%Y-%m-%d')} |
| 服务团队 | 甄云科技 |

---

## 统计摘要
| 指标 | 数值 |
|------|------|
| 总工单数 | {stats['total']} |
| SLA达标率 | {stats['sla']} |
| 产品满意度 | {stats['product_sat']} |
| 服务满意度 | {stats['service_sat']} |

---

## 第一部分：工单概览

### 1.1 按工单分类分布
{stats['category']}

### 1.2 按系统模块分布
{stats['module']}

### 1.3 月度分布
{stats['monthly']}

---

{result_0}

---

## 第二部分：分类维度深度分析

{result_1}

---

## 第三部分：模块维度深度分析

{result_2}

---

## 第四部分：高频Q&A知识库

{result_3}

---

## 第五部分：总结与改进建议

{result_4}

---

**报告结束**
"""

    # 报告完整性校验：必须包含五大部分的关键标记
    def _is_complete(filepath):
        markers = ["第一部分", "第二部分", "第三部分", "第四部分", "第五部分"]
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            return all(m in content for m in markers)
        except Exception:
            return False

    # 保存报告（跳过已存在的完整版本）
    version = 1
    output_file = os.path.join(OUTPUT_DIR, f"{CLIENT_NAME}_{REPORT_YEAR}_运维报告_V{version}.md")
    while os.path.exists(output_file) and not _is_complete(output_file):
        version += 1
        output_file = os.path.join(OUTPUT_DIR, f"{CLIENT_NAME}_{REPORT_YEAR}_运维报告_V{version}.md")
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(report)

    print(f"\n[OK] Report saved: {output_file}")
    return output_file


def convert_to_docx_and_send(md_file, client_name, year):
    """将MD转为DOCX并通过Feishu发送"""
    import subprocess, re
    from docx import Document

    docx_file = md_file.replace('.md', '.docx')

    # ---- MD → DOCX ----
    def add_table_from_markdown(doc, md_content):
        lines = md_content.strip().split('\n')
        table = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if re.match(r'\|[\s\-:|]+\|', line):
                continue
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if table is None:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    for i, cell in enumerate(cells):
                        table.rows[0].cells[i].text = cell
                else:
                    row = table.add_row()
                    for i, cell in enumerate(cells):
                        if i < len(row.cells):
                            row.cells[i].text = cell
        return table is not None

    def parse_and_convert(md_content, doc):
        lines = md_content.split('\n')
        i = 0
        table_buffer = []
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1; continue
            if line.startswith('|'):
                table_buffer.append(line)
                if i + 1 >= len(lines) or not lines[i+1].strip().startswith('|'):
                    add_table_from_markdown(doc, '\n'.join(table_buffer))
                    table_buffer = []
                i += 1; continue
            if line.startswith('# '):
                doc.add_heading(line[2:], 0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 3)
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                p = doc.add_paragraph()
                while '**' in text:
                    start = text.find('**'); end = text.find('**', start+2)
                    if end == -1: break
                    before = text[:start]; bold_text = text[start+2:end]; text = text[end+2:]
                    if before: p.add_run(before)
                    p.add_run(bold_text).bold = True
                if text: p.add_run(text)
            else:
                text = line
                p = doc.add_paragraph()
                while '**' in text:
                    start = text.find('**'); end = text.find('**', start+2)
                    if end == -1: break
                    before = text[:start]; bold_text = text[start+2:end]; text = text[end+2:]
                    if before: p.add_run(before)
                    p.add_run(bold_text).bold = True
                if text: p.add_run(text)
            i += 1

    print("  [3/3] MD → DOCX...")
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal'].font.size = 12
    parse_and_convert(md_content, doc)
    doc.save(docx_file)
    print(f"  [OK] DOCX saved: {docx_file}")

    # ---- 发送Feishu（先复制到白名单目录）----
    print("  [4/4] 发送Feishu...")
    feishu_target = "user:ou_de8266fa9b6ec7b8a25b58df4dab4e7f"
    # 飞书机器人只能发送白名单目录下的文件，先复制到 workspace/media
    media_basename = os.path.basename(docx_file)
    media_path = os.path.join(WORKSPACE_MEDIA_DIR, media_basename)
    import shutil
    shutil.copy2(docx_file, media_path)
    cmd = [
        "openclaw", "message", "send",
        "--channel", "feishu",
        "--target", feishu_target,
        "--media", media_path
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode == 0:
        print("  [OK] 已发送至Feishu")
    else:
        print(f"  [ERROR] Feishu发送失败: {result.stderr.strip()}")


def confirm_before_generation(client_name, year):
    """执行前确认，显示生成条件摘要"""
    import glob
    raw_dir = os.path.join(RAW_DATA_ROOT, client_name, "运维工单")
    files = sorted(glob.glob(os.path.join(raw_dir, "*.xlsx")))
    year_files = [f for f in files if str(year) in os.path.basename(f)]
    file_count = len(year_files)

    current_year = datetime.now().year
    default_note = f"（默认：当前{current_year}年，取上一自然年{year}）" if args.year is None else f"（指定）"

    print(f"\n{'='*50}")
    print(f"  报告生成确认")
    print(f"{'='*50}")
    print(f"  客户：{client_name}")
    print(f"  年份：{year}年 {default_note}")
    print(f"  依据：{raw_dir}")
    print(f"  找到：{file_count}个工单文件")
    if year_files:
        for f in year_files:
            print(f"    - {os.path.basename(f)}")
    print(f"{'='*50}")

    confirm = input(f"\n  确认继续生成？[Y/n]: ").strip().lower()
    if confirm == 'n':
        print("  已取消")
        sys.exit(0)
    print()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="SRM运维报告生成")
    parser.add_argument("client_name", help="客户名称")
    parser.add_argument("--year", type=int, default=None, help="年份（不指定则默认上一自然年）")
    args = parser.parse_args()
    year = args.year if args.year else datetime.now().year - 1
    build_paths(args.client_name, year)
    confirm_before_generation(args.client_name, year)
    md_file = generate_report()
    convert_to_docx_and_send(md_file, args.client_name, year)
